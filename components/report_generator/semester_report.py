import datetime
import time

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import docx2pdf

import config
import components.common.integrity as integrity
import components.common.metadata as metadata
import components.report_generator.document as document_helper
import components.report_generator.comment_generator as cgen
import components.report_generator.manifest as manifest
from components.common.grader_report import GraderReport

class Generator:
    """
    Report generator class for JARS data processor

    This class is responsible for generating a standardized DOCX file from a templated XLSX grader report.
    """

    def __init__(self, output_path, 
                 grader_report: GraderReport, 
                 date: datetime = None, 
                 signature_path = None, 
                 cgen_mode = "map", 
                 use_watermark = True):
        """
        Initialize the generator instance.

        Args:
            output_path (str): The path of the output file.
            grader_report (GraderReport): The grader report instance. Should be initialized before passing to this class.
            date (datetime): The date for the report.
            signature_path (str): The path to the signature image.
            cgen_mode (str): The comment generation mode. Can be "map" or "ai".
            use_watermark (bool): Whether to use a watermark in the generated reports.

        Returns:
            Generator: The report initialized generator instance.
        """

        time_now = datetime.datetime.now()
        init_time = time_now.strftime("%Y%m%d %H%M%S")

        print("[  ] Initializing generator...")
        
        self.output_path = output_path
        self.grader_report = grader_report
        self.date = date
        self.signature_path = signature_path
        self.cgen_mode = cgen_mode
        self.use_watermark = use_watermark
        if self.cgen_mode == "ai":
            self.manifest = manifest.Manifest(f"{self.output_path}/Manifest {init_time}.xlsx")

        print("[OK] Report generator initialized!")

    def generate_all(self, autocorrect = True, callback = None, force = False, convert_to_pdf = False, delay = False):
        """
        Generates reports for all students in the grader report.
        This function basically calls generate_for_student() for each student in the grader report.
        
        Args:
            autocorrect (bool): Whether to autocorrect the generated comments or not.
            callback (function): The callback function to be called after each student is processed.
            force (bool): Whether to force the generation process or not. This will ignore the data validation errors.
        """
        job_count = len(self.grader_report.students.index)
        
        job_start = datetime.datetime.now()
        print(f"[  ] Job started at {job_start.strftime('%Y-%m-%d %H:%M:%S')}. Generating reports for {job_count} students…")
        for i, student in enumerate(self.grader_report.students.index):
            status_message = f"Generating report for {student}…"
            print("\n========================")
            print(status_message)
            print("========================")
            if callback is not None:
                callback(i, job_count, status_message)
            
            start_time = datetime.datetime.now()
            self.generate_for_student(student_name = student, autocorrect = autocorrect, force = force)
            end_time = datetime.datetime.now()
            
            processing_time = end_time - start_time
            print(f"[OK] Report for {student} generated in {processing_time.total_seconds()} seconds.")
            print(f"Progress: {round(i / job_count * 100, 2)}%")

            if processing_time.total_seconds() < 4 and i < job_count - 1 and delay:
                wait_time = 4 - processing_time.total_seconds()
                print(f"[-'] Waiting for {wait_time} seconds before processing next student to avoid rate limiting…")
                time.sleep(wait_time)
        
        if callback is not None:
            callback(job_count, job_count, "")
            
        job_end = datetime.datetime.now()
        time_taken = job_end - job_start
        time_taken_formatted = f"{time_taken.seconds // 60} minutes {time_taken.seconds % 60} seconds"
        print(f"Progress: 100%")
        print(f"[OK] Job completed at {job_end.strftime('%Y-%m-%d %H:%M:%S')}. Time taken: {time_taken_formatted}")
        callback(job_count, job_count, f"Job completed at {job_end.strftime('%Y-%m-%d %H:%M:%S')}. Time taken: {time_taken_formatted}")

        if convert_to_pdf:
            docx2pdf.convert(self.output_path)
            for student in self.grader_report.students.index:
                if callback is not None:
                    callback(i, job_count, f"Creating PDF copy for {student}'s report…")
                print(f"[  ] Creating PDF copy for {student}'s report…")

                metadata.pdf_inject(f"{self.output_path}/{student}.pdf", student, self.grader_report)
                integrity.sign_pdf(f"{self.output_path}/{student}.pdf")
            
            if callback is not None:
                callback(i, job_count, f"PDF copies for all reports created!")
            print(f"[OK] PDF copies for all reports created!")

    def generate_for_student(self, student_name, autocorrect = True, force = False, convert_to_pdf = False):
        """
        Generates a report for a specific student.
        
        Args:
            student_name (str): The name of the student.
            autocorrect (bool): Whether to autocorrect the generated comments or not.
            force (bool): Whether to force the generation process or not. This will ignore the data validation errors.
        """
        # Validate data
        if not self.grader_report.data_valid and not force:
            print(f"[  ] Grader report incomplete. Aborting process...")
            return
        
        # Prepare data
        student_sna = {}
        str_letter_grade = str(self.grader_report.get_final_grade(student_name, "Letter Grade"))
        str_final_score = str(self.grader_report.get_final_grade(student_name, "Final Score"))
        count_sna = self.grader_report.count_sna()

        # Document processing begins
        document = Document()
        document = document_helper.setup_page(document, 'a4') # Setup page size to A4

        # Document metadata setup
        doc_prop = document.core_properties
        doc_prop.author = "JAC Academic Reporting System"
        doc_prop.title = f"{student_name} - {self.grader_report.get_course_info('Subject')} - S{self.grader_report.get_course_info('Semester')} AY{self.grader_report.get_course_info('School Year')} Report Card"
        doc_prop.subject = f"{self.grader_report.get_course_info('Subject')} - S{self.grader_report.get_course_info('Semester')} AY{self.grader_report.get_course_info('School Year')} Report Card"
        doc_prop.category = "Semester Report Card"
        doc_prop.revised = 1
        doc_prop.version = "1.0.0"
        doc_prop.keywords = "JAC; JARS; Report Card; Semester Report Card"
        doc_prop.language = "en-GB"
        doc_prop.content_status = "Final"

        # Spacing setup
        section_spacing = Pt(18)
        subject_description_spacing = Pt(12)

        if count_sna > 6:
            section_spacing = Pt(12)
        if count_sna > 8:
            section_spacing = Pt(10)
        if count_sna > 9:
            subject_description_spacing = Pt(8)
            section_spacing = Pt(8)
        
        # Sections and headers setup
        section = document.sections[0]
        header_content = section.header.paragraphs[0]
        header_content.alignment = WD_TABLE_ALIGNMENT.CENTER
        logo_run = header_content.add_run()
        logo_run.add_picture(config.get_config("logo_path"), width = Cm(5.56))
        
        # Default style setup (font, etc.)
        style = document.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        # CONTENT STARTS HERE
        # Top spacer
        top_spacer = document.add_paragraph()
        top_spacer.paragraph_format.space_before = Pt(0)
        top_spacer.paragraph_format.space_after = Pt(0)

        # Course Information Section
        ci_table = document.add_table(rows = 3, cols = 5)
        ci_table.style = "Table Grid"
        ci_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        ci_table.autofit = False
        ci_table.allow_autofit = False
        ci_table_col_widths = [Cm(3), Cm(8), Cm(3), Cm(1.5), Cm(1.5)]

        ci_table.cell(0, 0).text = "Student"
        ci_table.cell(0, 0).paragraphs[0].runs[0].bold = True
        ci_table.cell(0, 1).text = student_name

        ci_table.cell(1, 0).text = "Grade"
        ci_table.cell(1, 0).paragraphs[0].runs[0].bold = True
        ci_table.cell(1, 1).text = self.grader_report.get_course_info("Grade")
        
        ci_table.cell(2, 0).text = "School Year"
        ci_table.cell(2, 0).paragraphs[0].runs[0].bold = True
        ci_table.cell(2, 1).text = self.grader_report.get_course_info("School Year")

        ci_table.cell(0, 2).text = "Semester"
        ci_table.cell(0, 2).paragraphs[0].runs[0].bold = True
        ci_table.cell(0, 3).merge(ci_table.cell(0, 4)).text = str(self.grader_report.get_course_info("Semester"))
        ci_table.cell(0, 3).paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER

        ci_table.cell(1, 2).text = "Subject"
        ci_table.cell(1, 2).paragraphs[0].runs[0].bold = True
        ci_table.cell(1, 3).merge(ci_table.cell(1, 4)).text = self.grader_report.get_course_info("Subject")
        ci_table.cell(1, 3).paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
        
        ci_table.cell(2, 2).text = "Assessment"
        ci_table.cell(2, 2).paragraphs[0].runs[0].bold = True
        ci_table.cell(2, 3).text = str_final_score
        ci_table.cell(2, 3).paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
        ci_table.cell(2, 4).text =str_letter_grade
        ci_table.cell(2, 4).paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER

        for i in range(0, 3):
            for j in range(0, 5):
                ci_table.cell(i, j).width = ci_table_col_widths[j]

        # Subject Description Section
        sd_header = document.add_paragraph()
        sd_header.add_run("SUBJECT DESCRIPTION").bold = True
        sd_header.paragraph_format.space_before = subject_description_spacing
        sd_header.paragraph_format.space_after = Pt(0)
        
        sd_table = document.add_table(rows = 1, cols = 1)
        sd_table.style = "Table Grid"
        sd_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        sd_table.autofit = False
        sd_table.allow_autofit = False
        sd_table.rows[0].height = Cm(1.5)
        sd_table.cell(0, 0).text = self.grader_report.get_course_info("Subject Description")
        sd_table.cell(0, 0).paragraphs[0].alignment = WD_TABLE_ALIGNMENT.LEFT
        sd_table.cell(0, 0).width = Cm(17)

        # Skills and Assessment Section
        sna_header = document.add_paragraph()
        sna_header.add_run("SKILLS AND ASSESSMENT").bold = True
        sna_header.paragraph_format.space_before = section_spacing
        sna_header.paragraph_format.space_after = Pt(0)
        
        sna_table = document.add_table(rows = count_sna, cols = 2)
        sna_table.style = "Table Grid"
        sna_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        sna_table.autofit = False
        sna_table.allow_autofit = False
        
        for i, assessment in enumerate(self.grader_report.data_sna.columns):
            student_sna[assessment] = self.grader_report.get_grade_sna(student_name, assessment)
            sna_table.cell(i, 0).text = assessment
            sna_table.cell(i, 0).width = Cm(15)
            sna_table.cell(i, 1).text = str(student_sna[assessment])
            sna_table.cell(i, 1).width = Cm(2)
            sna_table.cell(i, 1).paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER

        # Personal Development Section
        pd_header = document.add_paragraph()
        pd_header.add_run("PERSONAL DEVELOPMENT").bold = True
        pd_header.paragraph_format.space_before = section_spacing
        pd_header.paragraph_format.space_after = Pt(0)
        
        pd_table = document.add_table(rows = len(self.grader_report.data_pd.columns) + 1, cols = 6)
        pd_table.style = "Table Grid"
        pd_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        pd_table.autofit = False
        pd_table.allow_autofit = False

        pd_table_header = pd_table.rows[0].cells
        pd_table_header[0].text = "Item"
        pd_table_header[1].text = "NI"
        pd_table_header[2].text = "S"
        pd_table_header[3].text = "G"
        pd_table_header[4].text = "VG"
        pd_table_header[5].text = "E"

        for i in range(0, 6): # Setup table header layout
            pd_table_header[i].paragraphs[0].runs[0].bold = True
            pd_table_header[i].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
            pd_table_header[i].width = Cm(1) if i != 0 else Cm(12)

        for i, item in enumerate(self.grader_report.data_pd.columns): # Setup table body (items)
            pd_table.cell(i + 1, 0).text = item
            pd_table.cell(i + 1, 0).paragraphs[0].alignment = WD_TABLE_ALIGNMENT.LEFT
            pd_table.cell(i + 1, 0).width = Cm(12)
            pd_grade = int(self.grader_report.get_grade_pd(student_name, item))
            pd_table.cell(i + 1, pd_grade).text = "✔"
            pd_table.cell(i + 1, pd_grade).paragraphs[0].runs[0].font.name = "Segoe UI Symbol"
            pd_table.cell(i + 1, pd_grade).paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
            
        for i in range(1, len(self.grader_report.data_pd.columns) + 1): # For each cell in row
            for j in range(1, 6):                                       # and in each column
                pd_table.cell(i, j).width = Cm(1)                       # set width to 1cm

        # Teacher's Comments Section        
        tc_header = document.add_paragraph()
        tc_header.add_run("TEACHER'S COMMENTS").bold = True
        tc_header.paragraph_format.space_before = section_spacing
        tc_header.paragraph_format.space_after = Pt(0)

        tc_table = document.add_table(rows = 1, cols = 1)
        tc_table.style = "Table Grid"
        tc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        tc_table.autofit = False
        tc_table.allow_autofit = False
        tc_table.rows[0].height = Cm(2)

        if self.cgen_mode == "map":
            comment_generator = cgen.CommentGenerator(student_name = student_name, 
                                                      short_name = self.grader_report.get_student_info(student_name, "Short Name"),
                                                      gender = self.grader_report.get_student_info(student_name, "Gender"),
                                                      comment_mapping = self.grader_report.data_comment_mapping,
                                                      student_result = student_sna,
                                                      letter_grade = str_letter_grade,
                                                     )
            tc_table.cell(0, 0).text = comment_generator.generate_comment(autocorrect = autocorrect)
        elif self.cgen_mode == "ai":
            comment_generator = cgen.AICommentGenerator(self.manifest)
            tc_table.cell(0, 0).text = comment_generator.generate_comment(nickname = self.grader_report.get_student_info(student_name, "Short Name"),
                                                                          gender = self.grader_report.get_student_info(student_name, "Gender"),
                                                                          final_grade = str_letter_grade,
                                                                          sna_list = student_sna,
                                                                          verbose = True
                                                                         )
            
        tc_table.cell(0, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        tc_table.cell(0, 0).width = Cm(17)

        # Acknowledgement Section
        ak_header = document.add_paragraph()
        ak_header.add_run("ACKNOWLEDGEMENT").bold = True
        ak_header.paragraph_format.space_before = section_spacing
        ak_header.paragraph_format.space_after = Pt(0)

        ak_table = document.add_table(rows = 2, cols = 4)
        ak_table.style = "Table Grid"
        ak_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        ak_table.autofit = False
        ak_table.allow_autofit = False

        for row in ak_table.rows:
            row.height = Cm(1.5) # Set row height to 1cm

        ak_table.cell(0, 0).paragraphs[0].add_run("Teacher:").bold = True
        ak_table.cell(0, 0).add_paragraph(self.grader_report.get_course_info("Teacher"))

        ak_table.cell(0, 1).text = "Signature"
        if self.signature_path is not None:
            ak_table.cell(0, 2).paragraphs[0].add_run().add_picture(self.signature_path, height = Cm(1.5))
        
        ak_table.cell(0, 3).text = "Date"
        if self.date is not None:
            ak_table.cell(0, 3).add_paragraph(self.date.strftime("%B %d, %Y"))
        
        ak_table.cell(0, 0).width = Cm(6.5)
        ak_table.cell(0, 1).width = Cm(2)
        ak_table.cell(0, 2).width = Cm(4.5)
        ak_table.cell(0, 3).width = Cm(4)
        document_helper.set_cell_border(ak_table.cell(0, 1), end = {"sz": 1, "val": "none"})
        document_helper.set_cell_border(ak_table.cell(0, 2), start = {"sz": 1, "val": "none"})

        ak_table.cell(1, 0).paragraphs[0].add_run("Parent:").bold = True
        ak_table.cell(1, 1).merge(ak_table.cell(1, 2)).text = "Signature"
        ak_table.cell(1, 3).text = "Date"
        ak_table.cell(1, 0).width = Cm(6.5)
        ak_table.cell(1, 1).width = Cm(6.5)
        ak_table.cell(1, 3).width = Cm(4)

        # Legend Section
        lg_header = document.add_paragraph()
        lg_header.add_run("GRADING SYSTEM").bold = True
        lg_header.paragraph_format.space_before = section_spacing
        lg_header.paragraph_format.space_after = Pt(0)

        lg_table = document.add_table(rows = 3, cols = 13)
        lg_table.style = "Table Grid"
        lg_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        lg_table.autofit = False
        lg_table.allow_autofit = False

        # Skills and Assessment legend header
        lg_table.cell(0, 0).merge(lg_table.cell(0, 5)).text = "Skills and Assessment"
        lg_table.cell(0, 0).paragraphs[0].runs[0].bold = True
        lg_table.cell(0, 0).width = Cm(8.4)

        # Small spacer between two legends
        lg_table.cell(0, 6).width = Cm(0.2)
        lg_table.cell(0, 6).text = ""
        document_helper.set_cell_border(lg_table.cell(0, 6), top = {"sz": 1, "val": "none"}, bottom = {"sz": 1, "val": "none"})
        document_helper.set_cell_margin(lg_table.cell(0, 6), margin_left_mm = 0, margin_right_mm = 0)
        
        lg_table.cell(1, 6).width = Cm(0.2)
        lg_table.cell(1, 6).text = ""
        document_helper.set_cell_border(lg_table.cell(1, 6), top = {"sz": 1, "val": "none"}, bottom = {"sz": 1, "val": "none"})
        document_helper.set_cell_margin(lg_table.cell(1, 6), margin_left_mm = 0, margin_right_mm = 0)
        
        lg_table.cell(2, 6).width = Cm(0.2)
        lg_table.cell(2, 6).text = ""
        document_helper.set_cell_border(lg_table.cell(2, 6), top = {"sz": 1, "val": "none"}, bottom = {"sz": 1, "val": "none"})
        document_helper.set_cell_margin(lg_table.cell(2, 6), margin_left_mm = 0, margin_right_mm = 0)

        # Personal Development legend header
        lg_table.cell(0, 7).merge(lg_table.cell(0, 12)).text = "Personal Development"
        lg_table.cell(0, 7).paragraphs[0].runs[0].bold = True
        lg_table.cell(0, 7).width = Cm(8.4)

        # Skills and Assessment legend
        lg_table.cell(1, 0).text = "A"
        lg_table.cell(1, 1).text = "95-100"
        lg_table.cell(2, 0).text = "B"
        lg_table.cell(2, 1).text = "85-94"
        lg_table.cell(1, 0).width = Cm(0.8)
        lg_table.cell(1, 1).width = Cm(1.8)
        lg_table.cell(2, 0).width = Cm(0.8)
        lg_table.cell(2, 1).width = Cm(1.8)
        document_helper.set_cell_margin(lg_table.cell(1, 0), margin_left_mm = 0, margin_right_mm = 0)
        document_helper.set_cell_margin(lg_table.cell(1, 1), margin_left_mm = 0, margin_right_mm = 0)
        document_helper.set_cell_margin(lg_table.cell(2, 0), margin_left_mm = 0, margin_right_mm = 0)
        document_helper.set_cell_margin(lg_table.cell(2, 1), margin_left_mm = 0, margin_right_mm = 0)

        lg_table.cell(1, 2).text = "C"
        lg_table.cell(1, 3).text = "75-84"
        lg_table.cell(2, 2).text = "D"
        lg_table.cell(2, 3).text = "40-74"
        lg_table.cell(1, 2).width = Cm(0.8)
        lg_table.cell(1, 3).width = Cm(1.8)
        lg_table.cell(2, 2).width = Cm(0.8)
        lg_table.cell(2, 3).width = Cm(1.8)

        lg_table.cell(1, 4).merge(lg_table.cell(2, 4)).text = "E"
        lg_table.cell(1, 5).merge(lg_table.cell(2, 5)).text = "Below 40"
        lg_table.cell(1, 4).width = Cm(0.8)
        lg_table.cell(1, 5).width = Cm(2)

        lg_table.cell(1, 7).text = "E"
        lg_table.cell(1, 8).text = "Excellent"
        lg_table.cell(2, 7).text = "VG"
        lg_table.cell(2, 8).text = "Very Good"
        lg_table.cell(1, 7).width = Cm(0.8)
        lg_table.cell(1, 8).width = Cm(1.8)
        lg_table.cell(2, 7).width = Cm(0.8)
        lg_table.cell(2, 8).width = Cm(1.8)

        lg_table.cell(1, 9).text = "G"
        lg_table.cell(1, 10).text = "Good"
        lg_table.cell(2, 9).text = "S"
        lg_table.cell(2, 10).text = "Satisfactory"
        lg_table.cell(1, 9).width = Cm(0.8)
        lg_table.cell(1, 10).width = Cm(2)
        lg_table.cell(2, 9).width = Cm(0.8)
        lg_table.cell(2, 10).width = Cm(2)

        lg_table.cell(1, 11).merge(lg_table.cell(2, 11)).text = "NI"
        lg_table.cell(1, 12).merge(lg_table.cell(2, 12)).text = "Needs Improvement"
        lg_table.cell(1, 11).width = Cm(0.8)
        lg_table.cell(1, 12).width = Cm(2.2)

        for i in range(0, 3): # Adjust cell formatting for legend table
            for j in range(0, 13):
                lg_table.cell(i, j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                lg_table.cell(i, j).paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
                lg_table.cell(i, j).paragraphs[0].runs[0].font.size = Pt(9)

        # Watermark setup
        if self.use_watermark:
            document_helper.add_image_watermark(document, config.get_config("watermark_path"), opacity = 0.1, width_pt = 600, height_pt = 600)

        # CONTENT ENDS HERE
        # Document processing ends
        # Save document. The output file will be named as the student's name.
        document.save(f"{self.output_path}/{student_name}.docx")
        time_docsaved = datetime.datetime.now()

        if convert_to_pdf:
            print(f"[  ] Creating a PDF copy for {student_name}'s report…")
            docx2pdf.convert(f"{self.output_path}/{student_name}.docx")
            metadata.pdf_inject(f"{self.output_path}/{student_name}.pdf", student_name, self.grader_report, time_docsaved)
            integrity.sign_pdf(f"{self.output_path}/{student_name}.pdf")
            print(f"[OK] PDF copy for {student_name}'s report created!")

        if self.cgen_mode == "ai":
            self.manifest.save()