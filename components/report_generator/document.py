import json

from docx import Document
from docx.shared import Mm, Pt
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

import config

@staticmethod
def setup_page(document: Document, page_size: str = 'a4'):
    """
    Sets up the page size of the document.
    
    Args:
        document (Document): The document to setup.
        page_size (str, optional): The page size to setup. Defaults to 'a4'.

    Returns:
        Document: The document with the page size setup.

    Raises:
        FileNotFoundError: When the page size preset file is not found.
    """
    page_param = json.load(open(f"{config.get_config('page_presets')}/{str(page_size)}.json"))
    
    section = document.sections[0]
    section.page_height = Mm(page_param['page_height'])
    section.page_width = Mm(page_param['page_width'])
    section.left_margin = Mm(page_param['margin_left'])
    section.right_margin = Mm(page_param['margin_right'])
    section.top_margin = Mm(page_param['margin_top'])
    section.bottom_margin = Mm(page_param['margin_bottom'])

    return document

@staticmethod
def set_cell_border(cell, **kwargs):
    """
    Set a table cell's border properties. Usage example:
    >>> set_cell_border(
    >>>     cell,
    >>>     top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
    >>>     bottom={"sz": 12, "color": "#00FF00", "val": "single"},
    >>>     start={"sz": 24, "val": "dashed", "shadow": "true"},
    >>>     end={"sz": 12, "val": "dashed"},
    >>> )

    Args:
        cell (docx.table._Cell): The cell to set the border of.
        **kwargs: The border properties to set.

    Raises:
        ValueError: When an invalid border property is passed.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

@staticmethod
def set_cell_margin(cell, margin_top_mm=0, margin_bottom_mm=0, margin_left_mm=0, margin_right_mm=0):
    """
    Sets the cell margins for a table.

    Args:
        cell (docx.table._Cell): The table cell to set the margins for.
        margin_top_mm (int, optional): The top margin in millimeters. Defaults to 0.
        margin_bottom_mm (int, optional): The bottom margin in millimeters. Defaults to 0.
        margin_left_mm (int, optional): The left margin in millimeters. Defaults to
        margin_right_mm (int, optional): The right margin in millimeters. Defaults to 0.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)

    margins = {
        'top': margin_top_mm,
        'bottom': margin_bottom_mm,
        'left': margin_left_mm,
        'right': margin_right_mm
    }

    for side, margin in margins.items():
        tag = 'w:{}'.format(side)
        element = tcMar.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tcMar.append(element)
        # Convert mm to twips (1 mm = 56.7 twips)
        element.set(qn('w:w'), str(int(margin * 56.7)))
        element.set(qn('w:type'), 'dxa')

@staticmethod
def add_image_watermark(doc, image_path, opacity=0.5, width_pt=350, height_pt=350):
    """
    Adds an image watermark to each page of the document.

    Args:
        doc (Document): The document to add the watermark to.
        image_path (str): The path to the image to use as the watermark.
        opacity (float, optional): The opacity of the watermark. Defaults to 0.5.
        width_pt (int, optional): The width of the watermark in points. Defaults to 350.
        height_pt (int, optional): The height of the watermark in points. Defaults to 350.
    """
    for section in doc.sections:
        header = section.header
        paragraph = header.paragraphs[0]

        run = paragraph.add_run()
        picture = run.add_picture(image_path, width=Pt(width_pt), height=Pt(height_pt))
        
        # Get inline element
        inline = picture._inline
        
        # Get parent of inline (the run element)
        parent = inline.getparent()
        
        # Convert to EMUs (English Metric Units) - Word's internal measurement
        cx = int(width_pt * 9525)
        cy = int(height_pt * 9525)
        
        # Build the anchor XML with proper namespaces
        opacity_val = int(opacity * 100000)
        
        # Get the rId for the image relationship
        blip = inline.graphic.graphicData.pic.blipFill.blip
        r_embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        
        anchor_xml = f'''
        <wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                   xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
                   xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"
                   xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
                   distT="0" distB="0" distL="0" distR="0" simplePos="0" 
                   relativeHeight="0" behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">
            <wp:simplePos x="0" y="0"/>
            <wp:positionH relativeFrom="margin">
                <wp:align>center</wp:align>
            </wp:positionH>
            <wp:positionV relativeFrom="margin">
                <wp:align>center</wp:align>
            </wp:positionV>
            <wp:extent cx="{cx}" cy="{cy}"/>
            <wp:effectExtent l="0" t="0" r="0" b="0"/>
            <wp:wrapNone/>
            <wp:docPr id="1" name="Watermark"/>
            <a:graphic>
                <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
                    <pic:pic>
                        <pic:nvPicPr>
                            <pic:cNvPr id="0" name="Watermark"/>
                            <pic:cNvPicPr/>
                        </pic:nvPicPr>
                        <pic:blipFill>
                            <a:blip r:embed="{r_embed}">
                                <a:alphaModFix amt="{opacity_val}"/>
                            </a:blip>
                            <a:stretch>
                                <a:fillRect/>
                            </a:stretch>
                        </pic:blipFill>
                        <pic:spPr>
                            <a:xfrm>
                                <a:off x="0" y="0"/>
                                <a:ext cx="{cx}" cy="{cy}"/>
                            </a:xfrm>
                            <a:prstGeom prst="rect">
                                <a:avLst/>
                            </a:prstGeom>
                        </pic:spPr>
                    </pic:pic>
                </a:graphicData>
            </a:graphic>
        </wp:anchor>
        '''
        
        # Parse the anchor XML
        anchor = parse_xml(anchor_xml)
        
        # Replace inline with anchor
        parent.replace(inline, anchor)
        
        # Center alignment
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER